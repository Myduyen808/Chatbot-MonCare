# ====== Fix for Windows ======
import sys
import os
import yaml
import re
import shutil
import hashlib
import json
from pathlib import Path
from datetime import datetime
import pandas as pd
if sys.platform == "win32":
    try: import mock
    except ImportError: os.system("pip install mock"); import mock
    sys.modules["pwd"] = mock.Mock()

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader, CSVLoader
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from docx import Document
from langchain_core.documents import Document as LC_Document
from pypdf import PdfReader

# Thêm cache ở đầu file, sau phần import
_vector_db_cache = None

def reset_vector_db_cache():
    """Xóa FAISS cache sau khi VectorDB thay đổi."""
    global _vector_db_cache
    _vector_db_cache = None

with open("db_config.yml", "r", encoding="utf-8") as f: db_config = yaml.safe_load(f)
with open("model_config.yml", "r", encoding="utf-8") as f: model_config = yaml.safe_load(f)


# =========================================================
# INACTIVE / SUPERSEDED SOURCES
# =========================================================

INACTIVE_SOURCES = {
    # Đã có hướng dẫn chính thức mới:
    # QĐ 318/QĐ-BYT năm 2026
    "Hướng dẫn đầy đủ về ăn dặm theo khuyến cáo của Viện Dinh dưỡng Quốc gia.docx":
        "superseded_by_QD_318_2026",

    # Tài liệu ngoài trọng tâm, từng gây nhiễu retrieval vitamin
    "20 trò chơi thú vị.docx":
        "out_of_scope",
}


def get_inactive_reason(source):

    filename = os.path.basename(
        str(source or "")
    ).strip().casefold()

    for name, reason in INACTIVE_SOURCES.items():

        if filename == name.casefold():
            return reason

    return None

# Thêm sau dòng _vector_db_cache = None
_embedding_model = None  
_source_metadata_cache = {}

def clean_pdf_text(text):
    """Xóa rác lỗi font thường gặp khi trích xuất PDF"""
    # 1. Xóa mã rác bullet point bị lỗi font (uf075, uf0b7...)
    text = re.sub(r'uf0[0-9a-fA-F]{2}', '', text)
    # 2. Thay chữ 'n' bị lỗi xuống dòng bằng khoảng trắng
    text = re.sub(r'\bn\b', ' ', text)
    # 3. Xóa các ký tự điều khiển ẩn
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    # 4. Xóa khoảng trắng thừa
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def load_embedding():
    global _embedding_model
    if _embedding_model is None:
        _embedding_model = HuggingFaceEmbeddings(model_name=model_config["embedding_path"])
    return _embedding_model

def clean_web_boilerplate(text):
    noise_keywords = [
        "Top of Form", "Bottom of Form", "ĐỂ LẠI THÔNG TIN TƯ VẤN", "Họ và tên*", 
        "Số điện thoại*", "Tôi đã đọc và đồng ý với Chính sách bảo vệ dữ liệu cá nhân", 
        "Cập nhật:", "Chia sẻ", "Tải và đặt lịch khám tự động trên ứng dụng MyVinmec",
        "Câu hỏi thảo luận", "anh/chị thấy gì", "nội dung tranh lật", 
        "hướng dẫn cách sử dụng tranh lật", "các kỹ năng cần thiết", "lời tranh",
        "để nấu một bữa", "Chatbot này tư vấn được những vấn đề gì"
    ]
    lines = text.split('\n')
    cleaned_lines = [line for line in lines if not any(noise in line for noise in noise_keywords)]
    clean_text = "\n".join(cleaned_lines).strip()
    paragraphs = clean_text.split('\n\n')
    unique_paragraphs, seen = [], set()
    for p in paragraphs:
        normalized_p = " ".join(p.split())
        if normalized_p not in seen and len(normalized_p) > 30: 
            seen.add(normalized_p); unique_paragraphs.append(p)
    return "\n\n".join(unique_paragraphs).strip()

def get_list_documents():
    documents = []

    file_configs = [
        ("word_path", ".docx"),
        ("pdf_path", ".pdf"),
        ("csv_path", ".csv"),
        ("excel_path", ".xlsx"),
    ]

    for folder_key, ext in file_configs:
        path = db_config.get(folder_key, "")

        if os.path.exists(path):
            for root, dirs, files in os.walk(path):
                for file in files:
                    if file.endswith(ext) and not file.startswith("~$"):
                        documents.append(file)

    return documents

def get_details(filename):
    text, filepath = "", ""

    if filename.endswith(".docx"):
        filepath = os.path.join(db_config["word_path"], filename)
        doc = Document(filepath)
        text = "\n".join([para.text for para in doc.paragraphs])

    elif filename.endswith(".pdf"):
        filepath = os.path.join(db_config["pdf_path"], filename)
        reader = PdfReader(filepath)
        for page in reader.pages:
            text += (page.extract_text() or "")

    elif filename.endswith(".csv"):
        filepath = os.path.join(
            db_config.get("csv_path", "data_store/csv"),
            filename
        )
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()

    elif filename.endswith(".xlsx"):
        filepath = os.path.join(
            db_config.get("excel_path", "data_store/excel"),
            filename
        )
        df = pd.read_excel(filepath)
        text = df.to_string(index=False)

    return {
        "Tên file": filename,
        "Đường dẫn": filepath,
        "Kích thước": os.path.getsize(filepath) if os.path.exists(filepath) else 0
    }, text

def get_document(filename):
    docs = []

    if filename.endswith(".docx"):
        doc = Document(os.path.join(db_config["word_path"], filename))
        docs.append(
            LC_Document(
                page_content="\n".join(
                    [p.text for p in doc.paragraphs if p.text.strip()]
                ),
                metadata={
                    "source": filename,
                    "file_type": "docx"
                }
            )
        )

    elif filename.endswith(".pdf"):
        docs = PyPDFLoader(
            os.path.join(db_config["pdf_path"], filename)
        ).load()

        for d in docs:
            d.page_content = clean_pdf_text(d.page_content)
            d.metadata["file_type"] = "pdf"

    elif filename.endswith(".csv"):
        docs = CSVLoader(
            os.path.join(
                db_config.get("csv_path", "data_store/csv"),
                filename
            ),
            encoding="utf-8"
        ).load()

        for d in docs:
            d.metadata["file_type"] = "csv"

    elif filename.endswith(".xlsx"):
        filepath = os.path.join(
            db_config.get("excel_path", "data_store/excel"),
            filename
        )

        docs = load_special_excel(filepath)

    return docs

def delete_document(filename):
    for path in [db_config["word_path"], db_config["pdf_path"], db_config.get("csv_path", "data_store/csv"),db_config.get("excel_path", "data_store/excel")]:
        filepath = os.path.join(path, filename)
        if os.path.exists(filepath): os.remove(filepath); return True
    return False

class MyPyWordLoader:
    def __init__(self, word_path=db_config["word_path"]): self.word_path = word_path
    def load(self):
        docx_documents = []
        for root, dirs, files in os.walk(self.word_path):
            for file in files:
                if file.endswith(".docx") and not file.startswith("~$"):
                    document = Document(os.path.join(root, file))
                    text = "\n".join([p.text for p in document.paragraphs if p.text.strip()])
                    if text.strip(): docx_documents.append(LC_Document(page_content=text, metadata={"source": file, "file_type": "docx"}))
        return docx_documents


def load_special_excel(file_path):
    """Nạp file Excel thành Document"""
    documents = []

    if file_path.endswith(".csv"):
        df = pd.read_csv(file_path)
    else:
        df = pd.read_excel(file_path)

    print(f"\nĐang đọc: {file_path}")
    print("Columns:", df.columns.tolist())
    print("Rows:", len(df))

    # chuẩn hóa tên cột
    df.columns = [str(c).strip().lower() for c in df.columns]

    TOPIC_MAP = {0: "Body Part", 1: "Disease", 2: "Drug", 3: "Medicine"}

    for _, row in df.iterrows():

        # ── Bộ đề sản khoa (medical_exam) ──────────────────────────
        if {"medical_topic", "question", "options", "answer"}.issubset(df.columns):

            content = (
                f"Chuyên khoa: {row.get('medical_topic', '')}. "
                f"Câu hỏi: {row.get('question', '')}. "
                f"Các lựa chọn: {row.get('options', '')}. "
                f"Đáp án: {row.get('answer', '')}."
            )

            metadata = {
                "source": os.path.basename(file_path),
                "file_type": "medical_exam"
            }

        # ── ViMedAQA format — CÓ CỘT context (Bo_De_Me_Va_Be.xlsx) ─
        # Ưu tiên trước FAQ vì có context đầy đủ hơn nhiều
        elif {"question", "answer", "context"}.issubset(df.columns):

            topic_id   = row.get("topic", "")
            topic_name = TOPIC_MAP.get(int(topic_id), "Y khoa") if str(topic_id).isdigit() else str(topic_id)
            title      = str(row.get("title", "")).strip()
            keyword    = str(row.get("keyword", "")).strip()
            question   = str(row.get("question", "")).strip()
            answer     = str(row.get("answer", "")).strip()
            context    = str(row.get("context", "")).strip()

            # Bỏ qua hàng rỗng
            if not question or not context:
                continue

            content = (
                f"Chủ đề: {topic_name}. "
                + (f"Tiêu đề: {title}. " if title and title != "nan" else "")
                + (f"Từ khoá: {keyword}. " if keyword and keyword != "nan" else "")
                + f"Câu hỏi: {question}. "
                + f"Trả lời: {answer}. "
                + f"Ngữ cảnh: {context}"
            )

            metadata = {
                "source"   : os.path.basename(file_path),
                "file_type": "vimedaqa",
                "topic"    : topic_name,
                "title"    : title,
            }

        # ── FAQ thông thường — chỉ có question + answer ─────────────
        elif {"question", "answer"}.issubset(df.columns):

            content = (
                f"Câu hỏi: {row.get('question', '')}. "
                f"Trả lời: {row.get('answer', '')}."
            )

            metadata = {
                "source": os.path.basename(file_path),
                "file_type": "faq"
            }

        else:
            print(
                f"⚠ Không nhận diện được cấu trúc file: "
                f"{os.path.basename(file_path)}"
            )
            continue

        documents.append(
            LC_Document(
                page_content=content,
                metadata=metadata
            )
        )

    return documents

# Thêm hàm này vào vectordb.py
def is_data_driven_chunk(text: str) -> bool:
    """
    Phát hiện đoạn văn chứa nhiều dữ liệu định lượng (bảng, thống kê).
    Nếu là bảng dữ liệu -> KHÔNG cắt nhỏ, giữ nguyên 1 chunk.
    """
    # Đếm số dòng chứa số liệu (mg, ml, tháng, tuổi, %)
    lines = text.split('\n')
    data_lines = 0
    for line in lines:
        if re.search(r'(\d+\s*(mg|ml|g|%|tháng|tuổi|ngày|lần|tuần|kcal|kg))', line.lower()):
            data_lines += 1
            
    # Nếu >30% dòng chứa số liệu -> đây là bảng/chuỗi thống kê -> Không cắt
    if len(lines) > 2 and (data_lines / len(lines)) > 0.3:
        return True
    return False

def _sha256_file(file_path):
    """Tính SHA-256 của một file."""

    sha256 = hashlib.sha256()

    with open(file_path, "rb") as file:

        for block in iter(
            lambda: file.read(1024 * 1024),
            b"",
        ):
            sha256.update(block)

    return sha256.hexdigest()

def _find_source_file(
    source,
    file_type,
):
    """Tìm đường dẫn thật của file nguồn."""

    source = str(
        source or ""
    ).strip()

    # Loader PDF/CSV đôi khi source đã là full path.
    if os.path.isfile(source):
        return os.path.abspath(source)

    filename = os.path.basename(
        source
    )

    folder_map = {
        "docx": db_config.get(
            "word_path",
            "",
        ),
        "pdf": db_config.get(
            "pdf_path",
            "",
        ),
        "csv": db_config.get(
            "csv_path",
            "",
        ),
        "faq": db_config.get(
            "excel_path",
            "",
        ),
        "vimedaqa": db_config.get(
            "excel_path",
            "",
        ),
        "medical_exam": db_config.get(
            "excel_path",
            "",
        ),
    }

    folder = folder_map.get(
        file_type,
        "",
    )

    if not folder:
        return ""

    candidate = os.path.join(
        folder,
        filename,
    )

    if os.path.isfile(candidate):
        return os.path.abspath(
            candidate
        )

    return ""

def add_management_metadata(
    metadata,
):
    """
    Thêm metadata quản trị nguồn:
    document_id, version_id, hash, updated_at.
    """

    metadata = dict(
        metadata or {}
    )

    source = metadata.get(
        "source",
        "",
    )

    file_type = metadata.get(
        "file_type",
        "",
    )

    filename = os.path.basename(
        str(source)
    )

    document_key = (
        f"{file_type}|"
        f"{filename.lower()}"
    )

    document_id = hashlib.sha256(
        document_key.encode("utf-8")
    ).hexdigest()[:16]

    source_path = _find_source_file(
        source,
        file_type,
    )

    cache_key = (
        source_path
        if source_path
        else document_key
    )

    if cache_key in _source_metadata_cache:

        metadata.update(
            _source_metadata_cache[
                cache_key
            ]
        )

        return metadata

    file_hash = ""
    updated_at = ""
    file_size = 0

    if (
        source_path
        and os.path.isfile(source_path)
    ):

        file_hash = _sha256_file(
            source_path
        )

        updated_at = datetime.fromtimestamp(
            os.path.getmtime(
                source_path
            )
        ).isoformat(
            timespec="seconds"
        )

        file_size = os.path.getsize(
            source_path
        )

    management_metadata = {
        "document_id": document_id,

        # 16 ký tự đầu của hash làm mã version.
        "version_id": (
            file_hash[:16]
            if file_hash
            else "unknown"
        ),

        "source_sha256": file_hash,

        "source_updated_at": updated_at,

        "source_size_bytes": file_size,
    }

    _source_metadata_cache[
        cache_key
    ] = management_metadata

    metadata.update(
        management_metadata
    )

    return metadata


def _write_vector_manifest(
    folder_path,
    chunk_count,
):
    """Tạo manifest kiểm tra tính toàn vẹn."""

    folder = Path(folder_path)

    index_faiss = folder / "index.faiss"
    index_pkl = folder / "index.pkl"

    if not index_faiss.exists():
        raise RuntimeError(
            "Không tạo được index.faiss"
        )

    if not index_pkl.exists():
        raise RuntimeError(
            "Không tạo được index.pkl"
        )

    manifest = {
        "built_at": datetime.now().isoformat(
            timespec="seconds"
        ),
        "chunk_count": int(chunk_count),
        "index_faiss_sha256": _sha256_file(
            index_faiss
        ),
        "index_pkl_sha256": _sha256_file(
            index_pkl
        ),
    }

    with open(
        folder / "manifest.json",
        "w",
        encoding="utf-8",
    ) as file:
        json.dump(
            manifest,
            file,
            ensure_ascii=False,
            indent=2,
        )

def _verify_vector_manifest(
    db_path,
):
    """Kiểm tra DB trước khi deserialize index.pkl."""

    folder = Path(db_path)

    manifest_path = (
        folder / "manifest.json"
    )

    if not manifest_path.exists():
        raise RuntimeError(
            "VectorDB chưa có manifest.json. "
            "Hãy rebuild VectorDB."
        )

    with open(
        manifest_path,
        "r",
        encoding="utf-8",
    ) as file:
        manifest = json.load(file)

    index_faiss = (
        folder / "index.faiss"
    )

    index_pkl = (
        folder / "index.pkl"
    )

    if not index_faiss.exists():
        raise RuntimeError(
            "Thiếu index.faiss"
        )

    if not index_pkl.exists():
        raise RuntimeError(
            "Thiếu index.pkl"
        )

    actual_faiss_hash = (
        _sha256_file(index_faiss)
    )

    actual_pkl_hash = (
        _sha256_file(index_pkl)
    )

    if (
        actual_faiss_hash
        != manifest.get(
            "index_faiss_sha256"
        )
    ):
        raise RuntimeError(
            "index.faiss bị thay đổi "
            "hoặc bị lỗi dữ liệu."
        )

    if (
        actual_pkl_hash
        != manifest.get(
            "index_pkl_sha256"
        )
    ):
        raise RuntimeError(
            "index.pkl bị thay đổi "
            "hoặc bị lỗi dữ liệu."
        )

    return manifest


def _safe_save_vector_db(
    db,
    db_path,
    chunk_count,
):
    """
    Build DB mới ở thư mục tạm.
    Chỉ thay DB cũ sau khi build thành công.
    """

    target_path = Path(db_path)

    temp_path = Path(
        str(db_path) + "_tmp"
    )

    backup_path = Path(
        str(db_path) + "_backup"
    )

    # Xóa thư mục tạm cũ nếu có.
    if temp_path.exists():
        shutil.rmtree(
            temp_path,
            ignore_errors=True,
        )

    if backup_path.exists():
        shutil.rmtree(
            backup_path,
            ignore_errors=True,
        )

    # 1. Save DB mới vào thư mục tạm.
    db.save_local(
        str(temp_path)
    )

    # 2. Tạo manifest.
    _write_vector_manifest(
        temp_path,
        chunk_count,
    )

    # 3. Kiểm tra artifact tối thiểu.
    if not (
        temp_path / "index.faiss"
    ).exists():
        raise RuntimeError(
            "VectorDB mới thiếu index.faiss"
        )

    if not (
        temp_path / "index.pkl"
    ).exists():
        raise RuntimeError(
            "VectorDB mới thiếu index.pkl"
        )

    # 4. Đổi tên DB cũ thành backup.
    if target_path.exists():
        target_path.rename(
            backup_path
        )

    try:
        # 5. Đưa DB mới vào vị trí chính.
        temp_path.rename(
            target_path
        )

    except Exception:

        # Nếu swap lỗi thì trả DB cũ trở lại.
        if (
            backup_path.exists()
            and not target_path.exists()
        ):
            backup_path.rename(
                target_path
            )

        raise

    # 6. Thành công thì bỏ backup.
    if backup_path.exists():
        shutil.rmtree(
            backup_path,
            ignore_errors=True,
        )

    reset_vector_db_cache()


def create_vectordb_with_file(pdf_path=db_config["pdf_path"], word_path=db_config["word_path"], csv_path=db_config.get("csv_path", "data_store/csv"), chunk_size=db_config["database_config"]["chunk_size"], chunk_overlap=db_config["database_config"]["chunk_overlap"], db_path=db_config["database_path"]):
    os.makedirs(pdf_path, exist_ok=True); os.makedirs(word_path, exist_ok=True); os.makedirs(csv_path, exist_ok=True)
    
    pdf_documents = DirectoryLoader(pdf_path, glob="*.pdf", loader_cls=PyPDFLoader).load()
    for doc in pdf_documents:
        doc.page_content = clean_pdf_text(doc.page_content) # <--- Làm sạch rác PDF
        doc.metadata["file_type"] = "pdf"
    
    word_documents = MyPyWordLoader(word_path).load()
    
    csv_documents_raw = DirectoryLoader(csv_path, glob="*.csv", loader_cls=CSVLoader, loader_kwargs={'encoding': 'utf-8'}).load()
    csv_documents = []
    for doc in csv_documents_raw:
        lines = [line.replace('noidung_doan:', '').strip() for line in doc.page_content.split('\n') if not (line.startswith('tieude:') or line.startswith('nguon:') or line.strip() == '')]
        clean_text = "\n".join(lines).strip()
        if len(clean_text) > 50: csv_documents.append(LC_Document(page_content=clean_text, metadata={**doc.metadata, "file_type": "csv"}))

    # Load Excel đặc biệt
    excel_path = db_config.get("excel_path", "data_store/excel")
    
    special_docs = []

    print("========== EXCEL DEBUG ==========")

    excel_path = db_config.get("excel_path", "data_store/excel")
    print("Excel path:", excel_path)
    print("Exists:", os.path.exists(excel_path))

    if os.path.exists(excel_path):

        for root, dirs, files in os.walk(excel_path):

            for file in files:

                if file.endswith(".xlsx") and not file.startswith("~$"):

                    full_path = os.path.join(root, file)

                    special_docs.extend(
                        load_special_excel(full_path)
                    )

                    print(f"Đã nạp file Excel: {file}")
                    print("Found Excel:", file)
            
    print("Số document Excel:", len(special_docs))

    documents = (
    pdf_documents
    + word_documents
    + csv_documents
    + special_docs
    )
    if not documents: print("Không tìm thấy tài liệu!"); return

    final_clean_documents = []

    for doc in documents:

        source = doc.metadata.get(
            "source",
            ""
        )

        inactive_reason = get_inactive_reason(
            source
        )

        if inactive_reason:

            print(
                "🗄️ [SOURCE EXCLUDED] "
                f"{source} | "
                f"reason={inactive_reason}"
            )

            continue

        cleaned_text = clean_web_boilerplate(
            doc.page_content
        )

        if len(cleaned_text) > 50:

            enriched_metadata = (
                add_management_metadata(
                    doc.metadata
                )
            )

            final_clean_documents.append(
                LC_Document(
                    page_content=cleaned_text,
                    metadata=enriched_metadata,
                )
            )
            
    # =========================================================
    # CHUNKING CÓ GIỚI HẠN KÍCH THƯỚC
    # =========================================================

    # Hard limit để tránh FAQ / ViMedAQA / bảng quá dài
    # lọt nguyên vẹn vào FAISS.
    MAX_CHUNK_CHARS = 1800

    # chunk_size trong RecursiveCharacterTextSplitter mặc định
    # được tính theo số ký tự, không phải token.
    effective_chunk_size = min(
        int(chunk_size),
        MAX_CHUNK_CHARS
    )

    # Không để overlap quá lớn so với chunk.
    effective_chunk_overlap = min(
        int(chunk_overlap),
        effective_chunk_size // 5
    )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=effective_chunk_size,
        chunk_overlap=effective_chunk_overlap,
    )

    final_chunks = []

    for doc in final_clean_documents:

        cleaned_text = clean_chunk_text(
            doc.page_content
        ).strip()

        if len(cleaned_text) < 50:
            continue

        file_type = doc.metadata.get(
            "file_type",
            ""
        )

        is_special = file_type in [
            "medical_exam",
            "faq",
            "vimedaqa",
        ]

        is_table = is_data_driven_chunk(
            cleaned_text
        )

        # -----------------------------------------------------
        # 1. Tài liệu đặc biệt hoặc bảng:
        #    - ngắn -> giữ nguyên
        #    - dài -> bắt buộc chia
        # -----------------------------------------------------
        if is_special or is_table:

            if is_table:
                chunk_type = "data_table"
                min_length = 80
            else:
                chunk_type = "special_text"
                min_length = 50

            # Ngắn -> giữ nguyên để bảo toàn cấu trúc
            if len(cleaned_text) <= effective_chunk_size:

                metadata = dict(doc.metadata)
                metadata["chunk_type"] = chunk_type

                final_chunks.append(
                    LC_Document(
                        page_content=cleaned_text,
                        metadata=metadata,
                    )
                )

            # Quá dài -> bắt buộc chia nhỏ
            else:

                metadata = dict(doc.metadata)
                metadata["chunk_type"] = chunk_type
                metadata["was_split"] = True

                temp_doc = LC_Document(
                    page_content=cleaned_text,
                    metadata=metadata,
                )

                chunks = splitter.split_documents(
                    [temp_doc]
                )

                for chunk in chunks:

                    chunk_text = (
                        clean_chunk_text(
                            chunk.page_content
                        ).strip()
                    )

                    if len(chunk_text) >= min_length:
                        final_chunks.append(
                            LC_Document(
                                page_content=chunk_text,
                                metadata=dict(
                                    chunk.metadata
                                ),
                            )
                        )

        # -----------------------------------------------------
        # 2. Tài liệu thông thường
        # -----------------------------------------------------
        else:

            metadata = dict(doc.metadata)
            metadata["chunk_type"] = "normal_text"

            temp_doc = LC_Document(
                page_content=cleaned_text,
                metadata=metadata,
            )

            chunks = splitter.split_documents(
                [temp_doc]
            )

            for chunk in chunks:

                chunk_text = (
                    clean_chunk_text(
                        chunk.page_content
                    ).strip()
                )

                if len(chunk_text) >= 50:
                    final_chunks.append(
                        LC_Document(
                            page_content=chunk_text,
                            metadata=dict(
                                chunk.metadata
                            ),
                        )
                    )
                    
    # Gắn mã chunk để phục vụ truy vết nguồn.
    for chunk_index, chunk in enumerate(
        final_chunks
    ):
        chunk.metadata[
            "chunk_id"
        ] = chunk_index

        raw_page = chunk.metadata.get(
            "page"
        )

        if isinstance(
            raw_page,
            int,
        ):
            chunk.metadata[
                "page_display"
            ] = raw_page + 1


    # DEBUG chỉ chạy MỘT LẦN sau vòng for.
    if final_chunks:

        chunk_lengths = [
            len(chunk.page_content)
            for chunk in final_chunks
        ]

        print(
            "\n========== CHUNK SIZE DEBUG =========="
        )

        print(
            "Effective chunk size:",
            effective_chunk_size,
        )

        print(
            "Effective overlap:",
            effective_chunk_overlap,
        )

        print(
            "Tổng chunks:",
            len(final_chunks),
        )

        print(
            "Chunk nhỏ nhất:",
            min(chunk_lengths),
        )

        print(
            "Chunk trung bình:",
            round(
                sum(chunk_lengths)
                / len(chunk_lengths),
                2,
            ),
        )

        print(
            "Chunk lớn nhất:",
            max(chunk_lengths),
        )

        oversized_chunks = [
            size
            for size in chunk_lengths
            if size > MAX_CHUNK_CHARS
        ]

        print(
            "Số chunk vượt hard limit:",
            len(oversized_chunks),
        )

        print(
            "======================================\n"
        )

    print(f"Chunks sau lọc: {len(final_chunks)}")

    db = FAISS.from_documents(
        documents=final_chunks,
        embedding=load_embedding()
    )

    _safe_save_vector_db(
    db=db,
    db_path=db_path,
    chunk_count=len(final_chunks),
    )

    print(f"Đã tạo FAISS DB với {len(final_chunks)} đoạn (Đã gắn thẻ file_type).")

def load_vector_db(
    db_path=db_config["database_path"]
):
    global _vector_db_cache

    if _vector_db_cache is None:

        # PHẢI kiểm tra trước khi deserialize pickle.
        manifest = _verify_vector_manifest(
            db_path
        )

        loaded_db = FAISS.load_local(
            db_path,
            load_embedding(),
            allow_dangerous_deserialization=True,
        )

        expected_chunks = int(
            manifest.get(
                "chunk_count",
                -1,
            )
        )

        actual_vectors = int(
            loaded_db.index.ntotal
        )

        actual_mappings = len(
            loaded_db.index_to_docstore_id
        )

        if expected_chunks >= 0:

            if (
                actual_vectors
                != expected_chunks
            ):
                raise RuntimeError(
                    "Số vector trong FAISS "
                    "không khớp manifest."
                )

            if (
                actual_mappings
                != expected_chunks
            ):
                raise RuntimeError(
                    "Docstore mapping "
                    "không khớp manifest."
                )

        _vector_db_cache = (
            loaded_db
        )

    return _vector_db_cache

def detect_query_priority(question):
    q = question.lower()
    mom_baby_keywords = ["núm vú", "cho con bú", "sữa mẹ", "sau sinh", "tắc sữa", "tắc tia sữa", "viêm tuyến vú", "mách sữa", "áp-xe vú", "massage vú", "hút sữa", "bầu vú", "rốn", "tưa miệng", "rôm sảy", "khóc dạ đề", "ăn dặm"]
    for kw in mom_baby_keywords:
        if kw in q: return ["docx", "pdf"]
    return ["pdf", "docx"]

# ===== SMART RETRIEVE ĐÃ SỬA =====
def smart_retrieve(question, llm, k=5, score_threshold=100.0):
    db = load_vector_db()

    # ── Tăng fetch_k để có pool lớn hơn, đặc biệt cho câu hỏi định lượng ──
    retriever = db.as_retriever(
        search_type="mmr",
        search_kwargs={
            "k": k,
            "fetch_k": 30,         # tăng từ 20 → 30
            "lambda_mult": 0.7     # tăng từ 0.6 → 0.7 (ưu tiên relevance hơn diversity)
        }
    )

    results = retriever.invoke(question)

    # ── Keyword overlap filter — dùng filtered thay vì bỏ qua nó ──
    question_words = set(question.lower().split())
    filtered = [
        doc for doc in results
        if len(question_words & set(doc.page_content.lower().split())) >= 1
    ]
    # fallback nếu không có doc nào vượt ngưỡng
    if not filtered:
        filtered = results

    all_results = []
    seen_contents = set()

    # ── Dùng filtered (đã fix bug: trước đây loop trên results, không phải filtered) ──
    for doc in filtered:
        content = str(doc.page_content)
        content_hash = hash(content[:200])
        
        if content_hash not in seen_contents:
            seen_contents.add(content_hash)
            all_results.append(doc)

    return all_results[:k]

# Thêm hàm này vào vectordb.py
def clean_chunk_text(text: str) -> str:
    """Xóa dòng rác nhưng GIỮ LẠI nội dung hữu ích trong cùng chunk"""
    noise_line_patterns = [
        "câu hỏi thảo luận",
        "anh/chị thấy gì",
        "nội dung tranh lật",
        "hướng dẫn cách sử dụng tranh lật",
        "các kỹ năng cần thiết",
        "lời tranh",
        "để nấu một bữa",   
    ]
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        line_lower = line.lower().strip()
        if any(p in line_lower for p in noise_line_patterns):
            continue
        if line_lower and line_lower[0].isdigit() and '.' in line_lower[:3]:
            continue
        cleaned.append(line)
    return '\n'.join(cleaned).strip()

def is_meaningful_chunk(text: str) -> bool:
    """Chỉ bỏ chunk nếu sau khi xóa noise vẫn còn quá ngắn"""
    cleaned = clean_chunk_text(text)
    return len(cleaned.strip()) >= 100

def get_retriever(k=5):
    return load_vector_db().as_retriever(search_kwargs={"k": k})

if __name__ == "__main__":
    print("Đang nạp dữ liệu..."); create_vectordb_with_file(); print("Xong!")